from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

LOGO_PATH = os.path.join(os.path.dirname(__file__), "static", "img", "zeonylogo.png")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: remove spacing before/after paragraph ────────────────────────────
def tight(para, space_before=0, space_after=60):
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)

# ═══════════════════════════════════════════════════════════════════════════════
#  HEADER ROW: logo right-aligned
# ═══════════════════════════════════════════════════════════════════════════════
# Use a 1-row, 2-col table: left=empty, right=logo
hdr_tbl = doc.add_table(rows=1, cols=2)
hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
# remove borders
for row in hdr_tbl.rows:
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top","left","bottom","right","insideH","insideV"):
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"),   "none")
            tag.set(qn("w:sz"),    "0")
            tag.set(qn("w:space"), "0")
            tag.set(qn("w:color"), "auto")
            tcBorders.append(tag)
        tcPr.append(tcBorders)

left_cell  = hdr_tbl.rows[0].cells[0]
right_cell = hdr_tbl.rows[0].cells[1]

# set column widths
left_cell.width  = Cm(11)
right_cell.width = Cm(6)

# logo in right cell
logo_para = right_cell.paragraphs[0]
logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
if os.path.exists(LOGO_PATH):
    run = logo_para.add_run()
    run.add_picture(LOGO_PATH, width=Cm(4))

# ── Title: QUOTATION ──────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
tight(title_para, space_before=6, space_after=10)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("QUOTATION")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

# thin horizontal rule via bottom border on a blank paragraph
def add_rule(document):
    hr = document.add_paragraph()
    tight(hr, 0, 4)
    pPr  = hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

add_rule(doc)

# ── Project Summary ───────────────────────────────────────────────────────────
ps = doc.add_paragraph()
tight(ps, space_before=8, space_after=6)
r1 = ps.add_run("Project summary")
r1.bold = True
r1.font.size = Pt(11)
r2 = ps.add_run(
    ": SGA Finance — Loan Management & Collection System built on Python (Flask), "
    "Supabase (PostgreSQL cloud database), and Bootstrap 5. Covers authentication & "
    "role-based access, customer onboarding with KYC documents, loan disbursement & "
    "tracking, payment collection by field agents, automatic daily compound penalty "
    "calculation, analytics dashboards, field user management, cloud file uploads, "
    "Vercel deployment, and production-ready scheduler."
)
r2.font.size = Pt(11)

# ── Section heading ───────────────────────────────────────────────────────────
sh = doc.add_paragraph()
tight(sh, space_before=10, space_after=4)
r = sh.add_run("Core fixed-price modules")
r.bold = True
r.font.size = Pt(11)

# ── Modules ───────────────────────────────────────────────────────────────────
modules = [
    ("01", "Project setup & core infrastructure",
     "Flask app factory, Supabase connectivity, environment config (.env), "
     "secret management, error handlers, session setup, Gunicorn production config.",
     "INR 15,000"),

    ("02", "Authentication & Authorization",
     "Login with email + password, Flask-Login session management, role-based "
     "access control (Admin / Field User), protected routes, logout.",
     "INR 20,000"),

    ("03", "User Management",
     "Admin CRUD for users: create, change password, delete. Role assignment "
     "(Admin / Field User). Password management UI with confirmation modal.",
     "INR 10,000"),

    ("04", "Customer Management & KYC",
     "Add / edit customers with full KYC: Aadhar, PAN, photo (mobile camera "
     "support), bank details. Auto-generated Customer IDs (CUS-YYYYMMDD-XXXX). "
     "Customer search and profile view.",
     "INR 25,000"),

    ("05", "Loan Creation & Lifecycle Management",
     "Create loans with auto-calculation of total interest, total repayable, daily EMI "
     "and due date. Live loan calculator (real-time preview). Loan ID auto-generation "
     "(LN-YYYYMMDD-XXXX). Status tracking: Active / Overdue / Cleared. "
     "Repayment progress bar.",
     "INR 30,000"),

    ("06", "Payment Collection Workflow",
     "Field user payment collection modal with notes. Balance before/after tracking. "
     "Auto-clear loan on full payment. Payment history per loan. "
     "Collected-by attribution.",
     "INR 18,000"),

    ("07", "Automatic Daily Penalty Engine",
     "APScheduler cron job (00:01 IST) calculating daily compound penalties on "
     "overdue outstanding balances. Penalty backfill for missed days. "
     "Penalty log per loan. Manual trigger button for admin. "
     "External scheduler HTTP hook (/hooks/run_penalty, Bearer token auth).",
     "INR 22,000"),

    ("08", "Field User Assignment & Dashboard",
     "Assign customers and loans to field agents. Field user dashboard: assigned loans, "
     "overdue count, collections today. Field user customer list (assigned only). "
     "Collect payment from customer profile.",
     "INR 18,000"),

    ("09", "Analytics Dashboard",
     "Admin KPI cards (disbursed, outstanding, collected, penalty). "
     "Monthly disbursement vs collections bar chart. Loan status doughnut chart. "
     "Monthly penalty trend line chart. Field user performance table. "
     "Top 10 overdue customers table.",
     "INR 22,000"),

    ("10", "File Uploads & Storage",
     "Customer photo and KYC document uploads (JPG/PNG/PDF, max 10 MB). "
     "Local disk storage with Supabase Storage fallback for Vercel/read-only environments. "
     "Filename sanitization and secure file serving.",
     "INR 12,000"),

    ("11", "REST API Endpoints",
     "POST /api/calculate_loan (real-time EMI preview). "
     "GET /api/loan_status/<id>. GET /api/customer_summary/<id>. "
     "All endpoints require authentication.",
     "INR 10,000"),

    ("12", "UI Polish & Responsive Design",
     "Bootstrap 5.3.2 responsive layout. Bootstrap Icons. Chart.js 4.4.2 charts. "
     "Color-coded status badges, progress bars, zebra tables. "
     "Mobile-friendly forms with camera capture.",
     "INR 12,000"),

    ("13", "Deployment, Documentation & Handover",
     "Vercel deployment support, Gunicorn server config, run.ps1 startup script. "
     "Environment variable documentation, Supabase schema SQL, README, "
     "IST timezone-aware scheduler, source code handover.",
     "INR 12,000"),
]

for num, title, desc, price in modules:
    p = doc.add_paragraph()
    tight(p, space_before=5, space_after=5)
    r_num = p.add_run(f"{num}. ")
    r_num.bold = True
    r_num.font.size = Pt(11)
    r_title = p.add_run(title)
    r_title.bold = True
    r_title.font.size = Pt(11)
    r_desc = p.add_run(f" ({desc})")
    r_desc.font.size = Pt(11)
    r_price = p.add_run(f": {price}")
    r_price.bold = True
    r_price.font.size = Pt(11)

add_rule(doc)

# ── Totals ────────────────────────────────────────────────────────────────────
total = sum([15000,20000,10000,25000,30000,18000,22000,18000,22000,12000,10000,12000,12000])

tp = doc.add_paragraph()
tight(tp, space_before=8, space_after=4)
tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = tp.add_run(f"Total (all modules): INR {total:,}")
r.bold = True
r.font.size = Pt(12)

gst_val = int(total * 0.18)
gst_p = doc.add_paragraph()
tight(gst_p, space_before=2, space_after=4)
gst_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r2 = gst_p.add_run(f"GST @ 18%: INR {gst_val:,}")
r2.font.size = Pt(11)

grand = total + gst_val
grand_p = doc.add_paragraph()
tight(grand_p, space_before=2, space_after=8)
grand_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r3 = grand_p.add_run(f"Grand Total (incl. GST): INR {grand:,}")
r3.bold = True
r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

add_rule(doc)

# ── Terms ─────────────────────────────────────────────────────────────────────
th = doc.add_paragraph()
tight(th, space_before=8, space_after=4)
r = th.add_run("Terms & Conditions")
r.bold = True
r.font.size = Pt(11)

terms = [
    "50% advance at order confirmation; balance 50% on delivery.",
    "Delivery timeline: 7 to 10 business days from advance receipt.",
    "GST @ 18% applicable on all amounts.",
    "Quotation valid for 30 days from April 30, 2026.",
    "Source code delivered after full payment is received.",
    "Support period begins from the date of final delivery.",
]
for t in terms:
    tp2 = doc.add_paragraph(style="List Bullet")
    tight(tp2, space_before=2, space_after=2)
    run = tp2.add_run(t)
    run.font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════════════════════
#  PAGE BREAK before tier section
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── WHAT IS BEING QUOTED ──────────────────────────────────────────────────────
wiq_h = doc.add_paragraph()
tight(wiq_h, space_before=0, space_after=4)
r = wiq_h.add_run("WHAT IS BEING QUOTED")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

wiq_p = doc.add_paragraph()
tight(wiq_p, space_before=2, space_after=8)
r = wiq_p.add_run(
    "A fully built, production-ready Loan Management & Collection Web Application built on "
    "Python (Flask), Supabase (PostgreSQL), and Bootstrap 5. The software includes customer "
    "onboarding, loan lifecycle management, payment collection, automated daily penalty "
    "calculations, analytics dashboards, and a field user mobile-friendly interface."
)
r.font.size = Pt(11)

add_rule(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  TIER COMPARISON TABLE
# ═══════════════════════════════════════════════════════════════════════════════
tc_h = doc.add_paragraph()
tight(tc_h, space_before=8, space_after=6)
r = tc_h.add_run("TIER COMPARISON")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

# Rows: (is_category, label, basic, pro, proplus)
tier_rows = [
    (True,  'PRICING', '', '', ''),
    (False, 'One-Time Cost',                              u'\u20b925,000',  u'\u20b945,000',  u'\u20b975,000'),
    (False, 'Annual Maintenance (AMC)',                   u'\u20b95,000/yr',u'\u20b98,000/yr',u'\u20b912,000/yr'),

    (True,  'CUSTOMER MANAGEMENT', '', '', ''),
    (False, 'Add / Edit Customers',                       'YES', 'YES', 'YES'),
    (False, 'KYC Upload (Aadhar, PAN, Photo)',            'YES', 'YES', 'YES'),
    (False, 'Bank Details Storage',                       'YES', 'YES', 'YES'),
    (False, 'Customer Search',                            'YES', 'YES', 'YES'),
    (False, 'Mobile Camera Photo Capture',                'YES', 'YES', 'YES'),
    (False, 'Customer ID Auto-Generation',                'YES', 'YES', 'YES'),

    (True,  'LOAN MANAGEMENT', '', '', ''),
    (False, 'Create Loans',                               'YES', 'YES', 'YES'),
    (False, 'Auto-Calculate Interest / EMI / Due Date',   'YES', 'YES', 'YES'),
    (False, 'Loan ID Auto-Generation',                    'YES', 'YES', 'YES'),
    (False, 'Loan Status Tracking (Active/Overdue/Cleared)', 'YES', 'YES', 'YES'),
    (False, 'Repayment Progress Bar',                     'YES', 'YES', 'YES'),
    (False, 'Loan Status Filters',                        'YES', 'YES', 'YES'),
    (False, 'Live Loan Calculator (real-time preview)',   'NO',  'YES', 'YES'),
    (False, 'Bulk Loan View with Days Overdue Badge',     'NO',  'YES', 'YES'),

    (True,  'PAYMENTS & PENALTIES', '', '', ''),
    (False, 'Record Payments',                            'YES', 'YES', 'YES'),
    (False, 'Payment History (Balance Before/After)',     'YES', 'YES', 'YES'),
    (False, 'Auto-Clear Loan on Full Payment',            'YES', 'YES', 'YES'),
    (False, 'Daily Penalty Auto-Calculation (APScheduler)','NO', 'YES', 'YES'),
    (False, 'Compound Daily Penalty on Outstanding',      'NO',  'YES', 'YES'),
    (False, 'Penalty Backfill (catches missed days)',     'NO',  'YES', 'YES'),
    (False, 'Penalty Log per Loan',                      'NO',  'YES', 'YES'),
    (False, 'Manual Penalty Trigger (Admin Button)',      'NO',  'YES', 'YES'),
    (False, 'External Scheduler Hook (/hooks/run_penalty)','NO', 'NO',  'YES'),

    (True,  'USER MANAGEMENT', '', '', ''),
    (False, 'Admin Role',                                 'YES', 'YES', 'YES'),
    (False, 'Field User Role',                            'NO',  'YES', 'YES'),
    (False, 'Create / Delete Users',                      'YES', 'YES', 'YES'),
    (False, 'Change Password',                            'YES', 'YES', 'YES'),
    (False, 'Assign Customers to Field Users',            'NO',  'YES', 'YES'),
    (False, 'Assign Loans to Field Users',                'NO',  'YES', 'YES'),
    (False, 'Field User Dashboard (My Loans / Collections Today)', 'NO', 'YES', 'YES'),
    (False, 'Field User Customer List (assigned only)',   'NO',  'YES', 'YES'),

    (True,  'ANALYTICS & REPORTING', '', '', ''),
    (False, 'Admin Dashboard KPI Cards',                  'YES', 'YES', 'YES'),
    (False, 'Recent Payments Table',                      'YES', 'YES', 'YES'),
    (False, 'Top Overdue Loans List',                     'NO',  'YES', 'YES'),
    (False, 'Monthly Disbursement vs Collections Bar Chart','NO','YES', 'YES'),
    (False, 'Loan Status Doughnut Chart',                 'NO',  'YES', 'YES'),
    (False, 'Monthly Penalty Trend Line Chart',           'NO',  'NO',  'YES'),
    (False, 'Field User Performance Table',               'NO',  'NO',  'YES'),
    (False, 'Top 10 Overdue Customers Table',             'NO',  'NO',  'YES'),

    (True,  'SYSTEM & DEPLOYMENT', '', '', ''),
    (False, 'Flask + Supabase Cloud Database',            'YES', 'YES', 'YES'),
    (False, 'Role-Based Access Control',                  'YES', 'YES', 'YES'),
    (False, 'Responsive UI (Bootstrap 5)',                'YES', 'YES', 'YES'),
    (False, 'Local File Upload (Static Folder)',          'YES', 'YES', 'YES'),
    (False, 'Supabase Storage (Cloud File Uploads)',      'NO',  'YES', 'YES'),
    (False, 'Vercel Deployment Support',                  'NO',  'YES', 'YES'),
    (False, 'Gunicorn Production Server Config',          'NO',  'YES', 'YES'),
    (False, 'API Endpoints (Calculate / Loan Status / Summary)', 'NO', 'YES', 'YES'),
    (False, 'IST Timezone-Aware Scheduler',               'NO',  'YES', 'YES'),

    (True,  'SUPPORT', '', '', ''),
    (False, 'Bug Fixes (post-delivery)',                  '30 days', '60 days', '90 days'),
    (False, 'Training Sessions',                          '1 session', '2 sessions', '3 sessions'),
    (False, 'Source Code Delivery',                       'YES', 'YES', 'YES'),
    (False, 'Deployment Assistance',                      'NO',  'YES', 'YES'),
    (False, 'Custom Branding (Logo, Colors)',             'NO',  'NO',  'YES'),
    (False, 'Priority Support',                           'NO',  'NO',  'YES'),
]

tbl = doc.add_table(rows=1 + len(tier_rows), cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# column widths
col_widths = [Cm(8.5), Cm(2.8), Cm(2.8), Cm(3.2)]
for ci, w in enumerate(col_widths):
    for cell in tbl.columns[ci].cells:
        cell.width = w

# header row
hdr_row = tbl.rows[0]
for ci, htxt in enumerate(['Feature', 'Basic', 'Pro', 'Pro Plus']):
    cell = hdr_row.cells[ci]
    cell.text = htxt
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cell.paragraphs[0].runs[0]
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(cell, '1A56DB')

# data rows
for ri, (is_cat, feat, basic, pro, proplus) in enumerate(tier_rows):
    row = tbl.rows[ri + 1]
    if is_cat:
        merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2]).merge(row.cells[3])
        merged.text = feat
        p = merged.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        shade_cell(merged, 'D6EAF8')
    else:
        for ci, val in enumerate([feat, basic, pro, proplus]):
            cell = row.cells[ci]
            cell.text = val
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.runs[0]
            r.font.size = Pt(9)
            if val == 'YES':
                r.font.color.rgb = RGBColor(0x1E, 0x8B, 0x3C)
                r.bold = True
            elif val == 'NO':
                r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            if ri % 2 == 0:
                shade_cell(cell, 'F8FBFF')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  PRICING SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
ps_h = doc.add_paragraph()
tight(ps_h, space_before=8, space_after=6)
r = ps_h.add_run("PRICING SUMMARY")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

ptbl = doc.add_table(rows=1, cols=4)
ptbl.style = 'Table Grid'
ptbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for ci, htxt in enumerate(['Tier', 'One-Time Cost', 'Annual Maintenance', 'Best For']):
    cell = ptbl.rows[0].cells[ci]
    cell.text = htxt
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cell.paragraphs[0].runs[0]
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(cell, '1A56DB')

pricing_data = [
    ('Basic',    u'\u20b925,000',  u'\u20b95,000/yr',  'Single admin, small operation, manual penalty tracking',    'EBF5FB'),
    ('Pro',      u'\u20b945,000',  u'\u20b98,000/yr',  'Small team with field agents, automated penalties, cloud hosting', 'D6EAF8'),
    ('Pro Plus', u'\u20b975,000',  u'\u20b912,000/yr', 'Full-scale NBFC / lending business with advanced analytics & priority support', 'AED6F1'),
]
for tier, ot, amc, best, clr in pricing_data:
    row = ptbl.add_row()
    for ci, val in enumerate([tier, ot, amc, best]):
        cell = row.cells[ci]
        cell.text = val
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if ci < 3 else WD_ALIGN_PARAGRAPH.LEFT
        if ci == 0:
            cell.paragraphs[0].runs[0].bold = True
        shade_cell(cell, clr)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  OPTIONAL ADD-ONS
# ═══════════════════════════════════════════════════════════════════════════════
ao_h = doc.add_paragraph()
tight(ao_h, space_before=8, space_after=6)
r = ao_h.add_run("OPTIONAL ADD-ONS")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

atbl = doc.add_table(rows=1, cols=2)
atbl.style = 'Table Grid'
atbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for ci, htxt in enumerate(['Add-On', 'Cost']):
    cell = atbl.rows[0].cells[ci]
    cell.text = htxt
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(cell, '1A56DB')

addons = [
    ('WhatsApp / SMS Payment Reminders Integration', u'\u20b910,000'),
    ('PDF Loan Agreement / Receipt Generation',       u'\u20b98,000'),
    ('Excel / CSV Export (Customers, Loans, Payments)',u'\u20b95,000'),
    ('Custom Mobile App (Android)',                   u'\u20b940,000'),
    ('Multi-Branch / Multi-Company Support',          u'\u20b920,000'),
    ('Additional Training Session (per session)',     u'\u20b92,000'),
]
for i, (addon, cost) in enumerate(addons):
    row = atbl.add_row()
    row.cells[0].text = addon
    row.cells[1].text = cost
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if i % 2 == 0:
        shade_cell(row.cells[0], 'F8FBFF')
        shade_cell(row.cells[1], 'F8FBFF')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  TERMS
# ═══════════════════════════════════════════════════════════════════════════════
add_rule(doc)

terms2_h = doc.add_paragraph()
tight(terms2_h, space_before=8, space_after=4)
r = terms2_h.add_run("TERMS")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

terms2 = [
    "50% advance at order confirmation, 50% on delivery.",
    "Delivery timeline: Basic (3-5 days), Pro (5-7 days), Pro Plus (7-10 days).",
    "GST applicable at 18% on all amounts.",
    "Prices valid for 30 days from date of quotation.",
]
for t in terms2:
    tp3 = doc.add_paragraph(style='List Bullet')
    tight(tp3, space_before=2, space_after=2)
    run = tp3.add_run(t)
    run.font.size = Pt(10)

note_p = doc.add_paragraph()
tight(note_p, space_before=8, space_after=4)
r = note_p.add_run(
    "Note: The software described above is already fully built. All tiers deliver the same "
    "codebase — tier pricing reflects licensing, deployment scope, support level, and "
    "feature enablement configuration."
)
r.font.size = Pt(10)
r.italic = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
tight(fp, 6, 0)
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("ZEONY TECHNOLOGIES  |  See the Future, Build it with Zeony  |  April 30, 2026")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
r.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), "SGA_Finance_Quotation.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
